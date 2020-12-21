# coding=utf-8           # -*- coding: utf-8 -*-
import requests
from bs4 import BeautifulSoup
from chardet import detect
from docx import Document
from os import listdir
from os.path import isfile, join

path = 'ii/'
text_file = "forecasts_doc.txt"

onlyfiles = [f for f in listdir(path) if isfile(join(path, f))]
all_text=""
for _fname in onlyfiles:
    # open doc
    document = Document( path+_fname )
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if len(para.text)>30:
                        all_text = all_text+para.text.lower()+"\n"
                        print(para.text)

    print(path+_fname)

with open(text_file, 'a') as f:
    f.write(all_text)

exit()

vgm_url   = 'http://meteoweb.ru/lfc.php'
html_text = requests.get(vgm_url).text
soup      = BeautifulSoup(html_text, 'html.parser')

base_url = 'http://meteoweb.ru/'

text_file = "forecasts.txt"

def download_text(count, item):
    if 'http' in item:
        return
        download_url = item
    else:
        download_url = '{}/{}'.format(base_url, item)

    _html_text   = requests.get(download_url).text
    soup         = BeautifulSoup(_html_text, 'html.parser')

    # Download the track
    print(download_url)
    for _ in soup.find_all('table',{'width':'800'}):
        print(_.text.encode("cp1251").decode('cp1251').encode('utf8') )
        # print( detect(_.text.encode('cp1251'))['encoding'] )
        # __ = str(_.text, 'cp1251').encode('utf-8')
        # print(__)
        exit()
        # r = requests.get(download_url, allow_redirects=True)
        with open(text_file, 'a') as f:
            f.write(_.text)
            print('=== Writed: {}'.format( download_url))

    # exit()
    # Print to the console to keep track of how the scraping is coming along.
    print('Downloaded: {}'.format( download_url))

count=0
for link in soup.find_all('a'):
    if '.php' in link.get('href'):
        print(link.get('href'))
        download_text(count,link.get('href'))
        count+=1
# print(soup)